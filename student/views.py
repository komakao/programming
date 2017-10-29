# -*- coding: UTF-8 -*-
#from django.shortcuts import render
from django.shortcuts import render_to_response, redirect
from django.contrib.auth.models import User
from django.http import HttpResponse
#from django.contrib.auth import authenticate, login
from django.template import RequestContext
#from django.contrib.auth.decorators import login_required
#from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.views.generic import ListView, CreateView
from django.core.exceptions import ObjectDoesNotExist
from django.utils.datastructures import MultiValueDictKeyError
from django.contrib.auth.models import Group
from teacher.models import Classroom
from student.models import Enroll, EnrollGroup, Work, Assistant, Exam, Bug, Debug, WorkFile
from account.models import Log, Message, MessagePoll, Profile, VisitorLog, Note
from show.models import Round
from certificate.models import Certificate
from student.forms import EnrollForm, GroupForm, SubmitForm, SeatForm, BugForm, DebugForm, DebugValueForm, GroupSizeForm
from django.utils import timezone
from student.lesson import *
from student.video import *
from account.avatar import *
from student.html2text import *
from account.models import Profile, PointHistory
from django.http import JsonResponse
from docx import *
from docx.shared import Inches
import StringIO
from datetime import datetime
from django.db.models import Q
from django.utils.timezone import localtime
import string
import base64
import sys
import os
import jieba
from django.conf import settings
from django.core.files.storage import FileSystemStorage
from uuid import uuid4
from wsgiref.util import FileWrapper
from itertools import groupby
from account.helper import VideoLogHelper
from django.contrib.auth.decorators import login_required

reload(sys)
sys.setdefaultencoding("utf-8")

# 判斷是否為授課教師
def is_teacher(user, classroom_id):
    return  user.groups.filter(name='teacher').exists() and Classroom.objects.filter(teacher_id=user.id, id=classroom_id).exists()

# 判斷是否開啟事件記錄
def is_event_open(request):
        enrolls = Enroll.objects.filter(student_id=request.user.id)
        for enroll in enrolls:
            classroom = Classroom.objects.get(id=enroll.classroom_id)
            if classroom.event_open:
                return True
        return False

# 判斷是否開啟課程事件記錄
def is_event_video_open(request):
        enrolls = Enroll.objects.filter(student_id=request.user.id)
        for enroll in enrolls:
            classroom = Classroom.objects.get(id=enroll.classroom_id)
            if classroom.event_video_open:
                return True
        return False

# 查看班級學生
def classmate(request, classroom_id):
        enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by("seat")
        enroll_group = []
        classroom_name=Classroom.objects.get(id=classroom_id).name
        for enroll in enrolls:
            login_times = len(VisitorLog.objects.filter(user_id=enroll.student_id))
            if enroll.group > 0 :
                enroll_group.append([enroll, EnrollGroup.objects.get(id=enroll.group).name, login_times])
            else :
                enroll_group.append([enroll, "沒有組別", login_times])
        # 記錄系統事件
        if is_event_open(request) :          
            log = Log(user_id=request.user.id, event=u'查看班級學生<'+classroom_name+'>')
            log.save()                        
        return render_to_response('student/classmate.html', {'classroom_name':classroom_name, 'enroll_group':enroll_group}, context_instance=RequestContext(request))

# 顯示所有組別
def group(request, classroom_id):
        student_groups = []
        classroom = Classroom.objects.get(id=classroom_id)
        group_open = Classroom.objects.get(id=classroom_id).group_open        
        groups = EnrollGroup.objects.filter(classroom_id=classroom_id)
        try:
                student_group = Enroll.objects.get(student_id=request.user.id, classroom_id=classroom_id).group
        except ObjectDoesNotExist :
                student_group = []		
        for group in groups:
            enrolls = Enroll.objects.filter(classroom_id=classroom_id, group=group.id)
            student_groups.append([group, enrolls, classroom.group_size-len(enrolls)])
            
        #找出尚未分組的學生
        def getKey(custom):
            return custom.seat	
        enrolls = Enroll.objects.filter(classroom_id=classroom_id)
        nogroup = []
        for enroll in enrolls:
            if enroll.group == 0 :
		        nogroup.append(enroll)		
	    nogroup = sorted(nogroup, key=getKey)

        # 記錄系統事件
        if is_event_open(request) :          
            log = Log(user_id=request.user.id, event=u'查看分組<'+classroom.name+'>')
            log.save()        
        return render_to_response('student/group.html', {'nogroup': nogroup, 'group_open': group_open, 'student_groups':student_groups, 'classroom':classroom, 'student_group':student_group, 'teacher': is_teacher(request.user, classroom_id)}, context_instance=RequestContext(request))

# 新增組別
def group_add(request, classroom_id):
        if request.method == 'POST':
            classroom_name = Classroom.objects.get(id=classroom_id).name            
            form = GroupForm(request.POST)
            if form.is_valid():
                group = EnrollGroup(name=form.cleaned_data['name'],classroom_id=int(classroom_id))
                group.save()
                
                # 記錄系統事
                if is_event_open(request) :                  
                    log = Log(user_id=request.user.id, event=u'新增分組<'+classroom_name+'><'+form.cleaned_data['name']+'>')
                    log.save()        
        
                return redirect('/student/group/'+classroom_id)
        else:
            form = GroupForm()
        return render_to_response('student/group_add.html', {'form':form}, context_instance=RequestContext(request))
        
# 設定組別人數
def group_size(request, classroom_id):
        if request.method == 'POST':
            form = GroupSizeForm(request.POST)
            if form.is_valid():
                classroom = Classroom.objects.get(id=classroom_id)
                classroom.group_size = form.cleaned_data['group_size']
                classroom.save()
                
                # 記錄系統事
                if is_event_open(request) :                  
                    log = Log(user_id=request.user.id, event=u'設定組別人數<'+classroom.name+'><'+str(form.cleaned_data['group_size'])+'>')
                    log.save()        
        
                return redirect('/student/group/'+classroom_id)
        else:
            classroom = Classroom.objects.get(id=classroom_id)
            form = GroupSizeForm(instance=classroom)
        return render_to_response('student/group_size.html', {'form':form}, context_instance=RequestContext(request))        

# 加入組別
def group_enroll(request, classroom_id,  group_id):
        classroom = Classroom.objects.get(id=classroom_id)
        members = Enroll.objects.filter(group=group_id)
        if len(members) < classroom.group_size:
            group_name = EnrollGroup.objects.get(id=group_id).name
            enroll = Enroll.objects.filter(student_id=request.user.id, classroom_id=classroom_id)
            enroll.update(group=group_id)
            # 記錄系統事件 
            if is_event_open(request) :          
                log = Log(user_id=request.user.id, event=u'加入組別<'+classroom.name+'><'+group_name+'>')
                log.save()         
        return redirect('/student/group/'+classroom_id)

# 刪除組別
def group_delete(request, group_id, classroom_id):
    group = EnrollGroup.objects.get(id=group_id)
    group.delete()
    classroom_name = Classroom.objects.get(id=classroom_id).name

    # 記錄系統事件 
    if is_event_open(request) :      
        log = Log(user_id=request.user.id, event=u'刪除組別<'+classroom_name+'><'+group.name+'>')
        log.save()       
    return redirect('/student/group/'+classroom_id)  
    
# 是否開放選組
def group_open(request, classroom_id, action):
    classroom = Classroom.objects.get(id=classroom_id)
    if action == "1":
        classroom.group_open=True
        classroom.save()
        # 記錄系統事件 
        if is_event_open(request) :          
            log = Log(user_id=request.user.id, event=u'開放選組<'+classroom.name+'>')
            log.save()            
    else :
        classroom.group_open=False
        classroom.save()
        # 記錄系統事件 
        if is_event_open(request) :          
            log = Log(user_id=request.user.id, event=u'關閉選組<'+classroom.name+'>')
            log.save()                
    return redirect('/student/group/'+classroom_id)  	
	
# 列出選修的班級
def classroom(request):
        classrooms = []
        enrolls = Enroll.objects.filter(student_id=request.user.id).order_by("-id")
        profile = Profile.objects.get(user_id=request.user.id)
        for enroll in enrolls:
            shows = Round.objects.filter(classroom_id=enroll.classroom_id)
            classrooms.append([enroll, shows])
        # 記錄系統事件 
        if is_event_open(request) :          
            log = Log(user_id=request.user.id, event='查看選修班級')
            log.save()          
        return render_to_response('student/classroom.html',{'classrooms': classrooms, 'profile':profile}, context_instance=RequestContext(request))    
    
# 查看可加入的班級
def classroom_add(request):
        classrooms = Classroom.objects.all().order_by('-id')
        classroom_teachers = []
        for classroom in classrooms:
            enroll = Enroll.objects.filter(student_id=request.user.id, classroom_id=classroom.id)
            if enroll.exists():
                classroom_teachers.append([classroom,classroom.teacher.first_name,1])
            else:
                classroom_teachers.append([classroom,classroom.teacher.first_name,0])   
        # 記錄系統事件 
        if is_event_open(request) :          
            log = Log(user_id=request.user.id, event='查看可加入的班級')
            log.save() 
        return render_to_response('student/classroom_add.html', {'classroom_teachers':classroom_teachers}, context_instance=RequestContext(request))
    
# 加入班級
def classroom_enroll(request, classroom_id):
        scores = []
        if request.method == 'POST':
                form = EnrollForm(request.POST)
                if form.is_valid():
                    try:
                        classroom = Classroom.objects.get(id=classroom_id)
                        if classroom.password == form.cleaned_data['password']:
                                enroll = Enroll(classroom_id=classroom_id, student_id=request.user.id, seat=form.cleaned_data['seat'])
                                enroll.save()
                                # 記錄系統事件 
                                if is_event_open(request) :  
                                    log = Log(user_id=request.user.id, event=u'加入班級<'+classroom.name+'>')
                                    log.save()                                 
                        else:
                                return render_to_response('message.html', {'message':"選課密碼錯誤"}, context_instance=RequestContext(request))
                      
                    except Classroom.DoesNotExist:
                        pass
                    
                    
                    return redirect("/student/group/" + str(classroom.id))
        else:
            form = EnrollForm()
        return render_to_response('student/classroom_enroll.html', {'form':form}, context_instance=RequestContext(request))
        
# 修改座號
def seat_edit(request, enroll_id, classroom_id):
    enroll = Enroll.objects.get(id=enroll_id)
    if request.method == 'POST':
        form = SeatForm(request.POST)
        if form.is_valid():
            enroll.seat =form.cleaned_data['seat']
            enroll.save()
            classroom_name = Classroom.objects.get(id=classroom_id).name
            # 記錄系統事件 
            if is_event_open(request) :              
                log = Log(user_id=request.user.id, event=u'修改座號<'+classroom_name+'>')
                log.save() 
            return redirect('/student/classroom')
    else:
        form = SeatForm(instance=enroll)

    return render_to_response('form.html',{'form': form}, context_instance=RequestContext(request))  

# 四種課程    
def lessons(request, unit):
        enrolls = Enroll.objects.filter(student_id=request.user.id)
        if request.user.is_authenticated():
            user_id = request.user.id
            user = User.objects.get(id=user_id)
            profile = Profile.objects.get(user=user)
            lock = profile.lock
        else :
            user_id = 0
            lock = 17
        # 記錄系統事件 
        if is_event_open(request):          
            log = Log(user_id=user_id, event=u'查看課程頁面<'+unit+'>')             
            log.save()         
        return render_to_response('student/lessons.html', {'unit': unit, 'lock':lock}, context_instance=RequestContext(request))

# 課程內容
def lesson(request, lesson):
        # 限登入者
        if not request.user.is_authenticated():
            return redirect("/account/login/")    
        else :
            lock = {'2':2, '3':3, '4':5, '5':7, '6':9, '7':11, '8':13, '9':14, '10':15, '11':16}
            # 記錄系統事件 
            # log = Log(user_id=request.user.id, event=u'課程內容<'+lesson+'>')
            #lo.save()
            # 改由 lesson_log 統一處理，含各課程之分頁 
            user = User.objects.get(id=request.user.id)
            profile = Profile.objects.get(user=user)
            if lesson in lock:
                if profile.lock < lock[lesson]:
                    if not user.groups.filter(name='teacher').exists():
                        return redirect("/")    
            classrooms = Classroom.objects.filter(teacher_id=request.user.id).order_by("-id")
            notes = Note.objects.filter(user_id=request.user.id, lesson = lesson).order_by('-id')
            return render_to_response('student/lesson.html', {'lesson':lesson, 'notes':notes, 'classrooms':classrooms}, context_instance=RequestContext(request))
        
def submit(request, lesson, index):
        scores = []
        workfiles = []
        works = Work.objects.filter(index=index, user_id=request.user.id)
        try:
            filepath = request.FILES['file']
        except :
            filepath = False
        if request.method == 'POST':
            if filepath :
                myfile = request.FILES['file']
                fs = FileSystemStorage()
                filename = uuid4().hex
                fs.save("static/work/"+str(request.user.id)+"/"+filename, myfile)
						
            form = SubmitForm(request.POST, request.FILES)

            if not works.exists():
                if form.is_valid():
                    work = Work(index=index, user_id=request.user.id, memo=form.cleaned_data['memo'], publication_date=timezone.now())
                    work.save()
                    workfile = WorkFile(work_id=work.id, filename=filename)
                    workfile.save()
										# credit
                    update_avatar(request.user.id, 1, 2)
                    # History
                    history = PointHistory(user_id=request.user.id, kind=1, message='2分--繳交作業<'+lesson_list[int(index)-1][2]+'>', url=request.get_full_path().replace("submit", "submitall"))
                    history.save()
                    # 記錄系統事件 
                    if is_event_open(request) :                      
                        log = Log(user_id=request.user.id, event=u'查看課程內容<'.encode("UTF-8")+lesson+u'> | 新增作業成功<'.encode("UTF-8")+index+'><'+lesson_list[int(index)-1][2]+'>')
                        log.save()    						
                    # lock + 1
                    user_id = request.user.id
                    user = User.objects.get(id=user_id)
                    profile = Profile.objects.get(user=user)
                    profile.lock += 1
                    profile.save()
            else:
                if form.is_valid():
                    #works[0].file = form.cleaned_data['file']
                    #works[0].memo = form.cleaned_data['memo']
                    #works[0].publication_date = timezone.localtime(timezone.now())
                    works.update(memo=form.cleaned_data['memo'],publication_date=timezone.localtime(timezone.now()))
                    workfile = WorkFile(work_id=works[0].id, filename=filename)
                    workfile.save()
                    # 記錄系統事件 
                    if is_event_open(request) :                      
                        log = Log(user_id=request.user.id, event=u'查看課程內容<'.encode("UTF-8")+lesson+u'> | 更新作業成功<'.encode("UTF-8")+index+'><'+lesson_list[int(index)-1][2]+'>')
                        log.save()                        
                else :
                    works.update(memo=form.cleaned_data['memo'])           
            return redirect('/student/submit/'+lesson+'/'+index)
        else:
            if not works.exists():
                form = SubmitForm()
            else:
                workfiles = WorkFile.objects.filter(work_id=works[0].id).order_by("-id")							
                form = SubmitForm(instance=works[0])
                if works[0].scorer>0: 
                    score_name = User.objects.get(id=works[0].scorer).first_name
                    scores = [works[0].score, score_name]
        lesson_name = lesson_list[int(index)-1]							
        return render_to_response('student/submit.html', {'form':form, 'scores':scores, 'index':index, 'lesson':lesson_name, 'workfiles': workfiles}, context_instance=RequestContext(request))

def work_download(request, index, user_id, workfile_id):
    workfile = WorkFile.objects.get(id=workfile_id)
    username = User.objects.get(id=user_id).first_name
    filename = username + "_" + lesson_list[int(index)-1][2]  + ".sb2"
    download =  settings.BASE_DIR + "/static/work/" + str(user_id) + "/" + workfile.filename
    wrapper = FileWrapper(file( download, "r" ))
    response = HttpResponse(wrapper, content_type = 'application/force-download')
    #response = HttpResponse(content_type='application/force-download')
    response['Content-Disposition'] = 'attachment; filename={0}'.format(filename.encode('utf8'))
    # It's usually a good idea to set the 'Content-Length' header too.
    # You can also set any other required headers: Cache-Control, etc.
    return response
    #return render_to_response('student/download.html', {'download':download})
			
# 列出所有作業        
def work(request, classroom_id):
    del lesson_list[:]
    reset()
    works = Work.objects.filter(user_id=request.user.id)
    for work in works:
        lesson_list[work.index-1].append(work.score)
        lesson_list[work.index-1].append(work.publication_date)
        if work.score > 0 :
            score_name = User.objects.get(id=work.scorer).first_name
            lesson_list[work.index-1].append(score_name)
        else :
            lesson_list[work.index-1].append("null")
    c = 0
    for lesson in lesson_list:
        assistant = Assistant.objects.filter(student_id=request.user.id, lesson=c+1, classroom_id=classroom_id)
        if assistant.exists() :
            lesson.append(1)
        else :
            lesson.append("")
        c = c + 1
        enroll_group = Enroll.objects.get(classroom_id=classroom_id, student_id=request.user.id).group
    # 記錄系統事件
    if is_event_open(request) :      
        log = Log(user_id=request.user.id, event=u'查看個人所有作業')
        log.save()          
    return render_to_response('student/work.html', {'works':works, 'lesson_list':lesson_list, 'user_id': request.user.id, 'classroom_id':classroom_id, 'group': enroll_group}, context_instance=RequestContext(request))

# 點擊各課tab記錄
def lesson_log(request, lesson):
    # 記錄系統事件
    tabname = request.POST.get('tabname')
    if is_event_open(request):   
        log = Log(user_id=request.user.id, event=u'查看課程內容<'+lesson+'> | '+tabname)
        log.save()
    elif is_event_video_open(request):
        if u"影片" in tabname:
            log = Log(user_id=request.user.id, event=u'查看課程內容<'+lesson+'> | '+tabname)
            log.save()
    return JsonResponse({'status':'ok'}, safe=False)


# 查詢某作業分組小老師    
def work_group(request, lesson, classroom_id):
        student_groups = []
        groups = EnrollGroup.objects.filter(classroom_id=classroom_id)
        try:
                student_group = Enroll.objects.get(student_id=request.user.id, classroom_id=classroom_id).group
        except ObjectDoesNotExist :
                student_group = []		
        for group in groups:
            enrolls = Enroll.objects.filter(classroom_id=classroom_id, group=group.id)
            group_assistants = []
            works = []
            scorer_name = ""
            for enroll in enrolls: 
                try:    
                    work = Work.objects.get(user_id=enroll.student_id, index=lesson)
                    if work.scorer > 0 :
                        scorer = User.objects.get(id=work.scorer)
                        scorer_name = scorer.first_name
                    else :
                        scorer_name = "X"
                except ObjectDoesNotExist:
                    work = Work(index=lesson, user_id=1)
                works.append([enroll, work.score, scorer_name, work.file])
                try :
                    assistant = Assistant.objects.get(student_id=enroll.student.id, classroom_id=classroom_id, lesson=lesson)
                    group_assistants.append(enroll)
                except ObjectDoesNotExist:
				    pass
            student_groups.append([group, works, group_assistants])
        lesson_data = lesson_list[int(lesson)-1]		
        # 記錄系統事件
        if is_event_open(request) :          
            log = Log(user_id=request.user.id, event=u'查看作業小老師<'+lesson+'>')
            log.save()        
        return render_to_response('student/work_group.html', {'lesson':lesson, 'lesson_data':lesson_data, 'student_groups':student_groups, 'classroom_id':classroom_id, 'student_group':student_group}, context_instance=RequestContext(request))

# 查詢某作業所有同學心得
def memo(request, classroom_id, index):
 
    enrolls = Enroll.objects.filter(classroom_id=classroom_id)
    datas = []
    for enroll in enrolls:
        try:
            work = Work.objects.get(index=index, user_id=enroll.student_id)
            datas.append([enroll, work.memo])
        except ObjectDoesNotExist:
            datas.append([enroll, ""])
    def getKey(custom):
        return custom[0].seat
    datas = sorted(datas, key=getKey)	
    # 記錄系統事件
    if is_event_open(request) :      
        log = Log(user_id=request.user.id, event=u'查看某作業所有同學心得<'+index+'>')
        log.save()    
    return render_to_response('student/memo.html', {'datas': datas}, context_instance=RequestContext(request))


# 查詢某班級心得
def memo_all(request, classroom_id):
        enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by("seat")
        classroom_name = Classroom.objects.get(id=classroom_id).name
        # 記錄系統事件
        if is_event_open(request) :          
            log = Log(user_id=request.user.id, event=u'查看班級心得<'+classroom_name+'>')
            log.save()            
        return render_to_response('student/memo_all.html', {'enrolls':enrolls, 'classroom_name':classroom_name}, context_instance=RequestContext(request))

# 查詢某班級心得統計
def memo_count(request, classroom_id):
        enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by("seat")
        members = []
        for enroll in enrolls:
            members.append(enroll.student_id)
        classroom = Classroom.objects.get(id=classroom_id)
        works = Work.objects.filter(user_id__in=members)
        memo = ""
        for work in works:
            memo += work.memo
        memo = memo.rstrip('\r\n')
        seglist = jieba.cut(memo, cut_all=False)
        hash = {}
        for item in seglist: 
            if item in hash:
                hash[item] += 1
            else:
                hash[item] = 1
        words = []
        count = 0
        error=""
        for key, value in sorted(hash.items(), key=lambda x: x[1], reverse=True):
            if ord(key[0]) > 32 :
                count += 1	
                words.append([key, value])
                if count == 30:
                    break
        # 記錄系統事件
        if is_event_open(request) :          
            log = Log(user_id=request.user.id, event=u'查看班級心得統計<'+classroom.name+'>')
            log.save()            
        return render_to_response('student/memo_count.html', {'words':words, 'enrolls':enrolls, 'classroom':classroom}, context_instance=RequestContext(request))

# 評分某同學某進度心得
@login_required
def memo_user(request, user_id):
    user = User.objects.get(id=user_id)
    del lesson_list[:]
    reset()
    works = Work.objects.filter(user_id=user_id)
    for work in works:
        lesson_list[work.index-1].append(work.memo)

    # 記錄系統事件
    if is_event_open(request) :        
        log = Log(user_id=request.user.id, event=u'查閱個人心得<'+user.first_name+'>')
        log.save()  
    return render_to_response('student/memo_user.html', {'lesson_list':lesson_list, 'student': user}, context_instance=RequestContext(request))


# 查詢某班級某作業心得統計
def memo_work_count(request, classroom_id, work_id):
        enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by("seat")
        members = []
        for enroll in enrolls:
            members.append(enroll.student_id)
        classroom = Classroom.objects.get(id=classroom_id)
        works = Work.objects.filter(user_id__in=members, index=int(work_id))
        memo = ""
        for work in works:
            memo += work.memo
        memo = memo.rstrip('\r\n')
        seglist = jieba.cut(memo, cut_all=False)
        hash = {}
        for item in seglist: 
            if item in hash:
                hash[item] += 1
            else:
                hash[item] = 1
        words = []
        count = 0
        for key, value in sorted(hash.items(), key=lambda x: x[1], reverse=True):
            count += 1
            if ord(key[0]) > 32 :
                words.append([key, value])
                if count == 30:
                    break
        # 記錄系統事件
        if is_event_open(request) :          
            log = Log(user_id=request.user.id, event=u'查看班級作業心得統計<'+classroom.name+'><'+work_id+'>')
            log.save()                    
        return render_to_response('student/memo_work_count.html', {'words':words, 'enrolls':enrolls, 'classroom':classroom,  'work_id':work_id, 'lesson':lesson_list[int(work_id)-1][2]}, context_instance=RequestContext(request))

			
# 查詢某班某詞句心得
def memo_word(request, classroom_id, word):
        enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by("seat")
        members = []
        for enroll in enrolls:
            members.append(enroll.student_id)
        classroom = Classroom.objects.get(id=classroom_id)
        works = Work.objects.filter(user_id__in=members, memo__contains=word).order_by('index')
        for work in works:
            work.memo = work.memo.replace(word, '<font color=red>'+word+'</font>')
        # 記錄系統事件
        if is_event_open(request) :          
            log = Log(user_id=request.user.id, event=u'查看班級心得詞句<'+classroom.name+'><'+word+'>')
            log.save()            
        return render_to_response('student/memo_word.html', {'word':word, 'works':works, 'classroom':classroom}, context_instance=RequestContext(request))
		
# 查詢某班某作業某詞句心得
def memo_work_word(request, classroom_id, work_id, word):
        enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by("seat")
        members = []
        for enroll in enrolls:
            members.append(enroll.student_id)
        classroom = Classroom.objects.get(id=classroom_id)
        works = Work.objects.filter(user_id__in=members, index=work_id, memo__contains=word)
        for work in works:
            work.memo = work.memo.replace(word, '<font color=red>'+word+'</font>')
        # 記錄系統事件
        if is_event_open(request) :          
            log = Log(user_id=request.user.id, event=u'查看班級心得詞句<'+classroom.name+'><'+word+'>')
            log.save()            
        return render_to_response('student/memo_work_word.html', {'word':word, 'works':works, 'classroom':classroom, 'lesson':lesson_list[int(work_id)-1][2]}, context_instance=RequestContext(request))
		
		
# 查詢個人心得
def memo_show(request, user_id, unit,classroom_id, score):
    user_name = User.objects.get(id=user_id).first_name
    del lesson_list[:]
    reset()
    works = Work.objects.filter(user_id=user_id)
    for work in works:
        lesson_list[work.index-1].append(work.score)
        lesson_list[work.index-1].append(work.publication_date)
        if work.score > 0 :
            score_name = User.objects.get(id=work.scorer).first_name
            lesson_list[work.index-1].append(score_name)
        else :
            lesson_list[work.index-1].append("null")
        lesson_list[work.index-1].append(work.memo)
    c = 0
    for lesson in lesson_list:
        assistant = Assistant.objects.filter(student_id=user_id, lesson=c+1)
        if assistant.exists() :
            lesson.append("V")
        else :
            lesson.append("")
        c = c + 1
        #enroll_group = Enroll.objects.get(classroom_id=classroom_id, student_id=request.user.id).group
    user = User.objects.get(id=user_id)
    # 記錄系統事件
    if is_event_open(request) :      
        log = Log(user_id=request.user.id, event=u'查看同學心得<'+user_name+'><'+unit+'>')
        log.save()        
    return render_to_response('student/memo_show.html', {'classroom_id': classroom_id, 'works':works, 'lesson_list':lesson_list, 'user_name': user_name, 'unit':unit, 'score':score}, context_instance=RequestContext(request))

# 查詢作業進度
def progress(request, classroom_id, unit):
    bars = []
    bars1 = []
    bars2 = []
    bars3 = []
    bars4 = []
    a = 0
    classroom = Classroom.objects.get(id=classroom_id)
    enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by("seat")
    for enroll in enrolls:
        c = 0
        for lesson in lesson_list :
            works = Work.objects.filter(user_id=enroll.student_id)
            bars.append([enroll, [], ""])
            for work in works:
                if work.index == c+1:
                    bars[a*41+c][1] = work
                    if work.scorer > 0 :
                        score_name = User.objects.get(id=work.scorer).first_name
                        bars[a*41+c][2] = score_name
            c = c + 1
        for i in range(17) :
            bars1.append(bars[i+41*a])
        for i in range(8) :
            bars2.append(bars[i+17+41*a])
        for i in range(8) :
            bars3.append(bars[i+25+41*a])
        for i in range(8) :
            bars4.append(bars[i+33+41*a])
        a = a + 1
    # 記錄系統事件
    if is_event_open(request) :      
        log = Log(user_id=request.user.id, event=u'查看作業進度<'+unit+'><'+classroom.name+'>')
        log.save()           
    return render_to_response('student/progress.html', {'unit':unit, 'bars1':bars1, 'bars2':bars2, 'bars3':bars3, 'bars4':bars4,'classroom':classroom, 'lesson_list': lesson_list,}, context_instance=RequestContext(request))
    

# 積分排行榜
class RankListView(ListView):
    context_object_name = 'datas'
    #paginate_by = 100
    template_name = 'student/rank_list.html'
    def get_queryset(self):
        datas = []
        enrolls = Enroll.objects.filter(classroom_id=self.kwargs['classroom_id'])
        for enroll in enrolls:
            try:
                profile = Profile.objects.get(user_id=enroll.student_id)
                if self.kwargs['kind'] == "0":
                    value = profile.work + profile.assistant + profile.debug + profile.creative
                    datas.append([enroll, value])
                elif self.kwargs['kind'] == "1":
                    datas.append([enroll, profile.work])
                elif self.kwargs['kind'] == "2":
                    datas.append([enroll, profile.assistant])
                elif self.kwargs['kind'] == "3":
                    datas.append([enroll, profile.debug])
                elif self.kwargs['kind'] == "4":
                    datas.append([enroll, profile.creative])
            except ObjectDoesNotExist:
                pass
        def getKey(custom):
            return custom[1], custom[0].seat	
        datas = sorted(datas, key=getKey, reverse=True)		
        # 記錄系統事件
        if is_event_open(self.request) :          
            log = Log(user_id=self.request.user.id, event=u'查看積分排行榜<'+self.kwargs['kind']+'>')
            log.save()          
        return datas
		
    def get_context_data(self, **kwargs):
        context = super(RankListView, self).get_context_data(**kwargs)
        context['kind'] = self.kwargs['kind']
        return context	

# 測驗卷
def exam(request):
    # 限登入者
    if not request.user.is_authenticated():
        return redirect("/account/login/")    
    else :
        return render_to_response('student/exam.html', context_instance=RequestContext(request))

# 測驗卷得分
def exam_score(request):
        exams = Exam.objects.filter(student_id=request.user.id)
        return render_to_response('student/exam_score.html', {'exams':exams} , context_instance=RequestContext(request))

# 測驗卷檢查答案		
def exam_check(request):
    exam_id = request.POST.get('examid')
    user_answer = request.POST.get('answer').split(",")
    # 記錄系統事件
    if is_event_open(request) :      
        log = Log(user_id=request.user.id, event=u'繳交測驗卷<'+exam_id+'>')
        log.save()      
    if exam_id == "1":
        answer = "C,A,D,C,C,A,B,B,D,D"
        answer_list = answer.split(",")
        ''' 儲存答案 '''
        ua_test = ""
        score = 0
        for i in range(10) :
            if user_answer[i] == answer_list[i] :
                score = score + 10
            i = i + 1
            ua_test = "".join(user_answer)
            '''ua_test = ua_test + user_answer[i]
            '''
        exam = Exam(exam_id=1, student_id=request.user.id, answer=ua_test, score=score)
        exam.save()
        ''' 回傳正確答案 '''
        return JsonResponse({'status':'ok','answer':answer}, safe=False)	
    elif exam_id == "2":
        answer = "B,C,C,A,D,B,A,D,B,C"
        answer_list = answer.split(",")
        ''' 儲存答案 '''
        ua_test = ""
        score = 0
        for i in range(10) :
            if user_answer[i] == answer_list[i] :
                score = score + 10
            i = i + 1
            ua_test = "".join(user_answer)
            '''ua_test = ua_test + user_answer[i]
            '''
        exam = Exam(exam_id=2, student_id=request.user.id, answer=ua_test, score=score)
        exam.save()	
        return JsonResponse({'status':'ok','answer':answer}, safe=False)
    elif exam_id == "3":
        answer = "D,C,A,B,D,C,D,A,D,B"
        answer_list = answer.split(",")
        ''' 儲存答案 '''
        ua_test = ""
        score = 0
        for i in range(10) :
            if user_answer[i] == answer_list[i] :
                score = score + 10
            i = i + 1
            ua_test = "".join(user_answer)
            '''ua_test = ua_test + user_answer[i]
            '''
        exam = Exam(exam_id=3, student_id=request.user.id, answer=ua_test, score=score)
        exam.save()	
        return JsonResponse({'status':'ok','answer':answer}, safe=False)
    else:
        return JsonResponse({'status':'ko'}, safe=False)
        
def debug_value(request, bug_id):
        import sys
        reload(sys)
        sys.setdefaultencoding("utf-8")
        # A comment was posted
        debug = Debug.objects.filter(id=3)
        debug_value_form = DebugValueForm(data=request.POST)		
        if debug_value_form.is_valid():
            try:
                debug_id = int(debug_value_form.cleaned_data['id']) 
                debug = Debug.objects.get(id=debug_id)
                url = "/student/bug/"+str(bug_id)
                if debug_value_form.cleaned_data['reward'] == "0":
                    reward = "沒有解決"
                    msg =u'0分<'+  request.user.first_name.encode('utf-8')+u'>評價您的除錯<'+reward.encode('utf-8')+'>'
                elif debug_value_form.cleaned_data['reward'] == "1":
                    reward = "部份解決"
                    msg =u'1分<'+  request.user.first_name.encode('utf-8')+u'>評價您的除錯<'+reward.encode('utf-8')+'>'
                elif debug_value_form.cleaned_data['reward'] == "2":
                    reward = "大概解決"
                    msg =u'2分<'+  request.user.first_name.encode('utf-8')+u'>評價您的除錯<'+reward.encode('utf-8')+'>'
                elif debug_value_form.cleaned_data['reward'] == "3":
                    reward = "完全解決"						
                    msg =u'3分<'+ request.user.first_name.encode('utf-8')+u'>評價您的除錯<'+reward.encode('utf-8')+'>'


				# credit
                if debug.reward < 0 :             
					# credit
                    update_avatar(debug.author_id, 3, int(debug_value_form.cleaned_data['reward']))
                    # History
                    history = PointHistory(user_id=debug.author_id, kind=3, message=msg, url=url)
                    history.save()	

                # create Message
                message = Message.create(title=msg, url=url, time=timezone.localtime(timezone.now()))
                message.save()
                messagepoll = MessagePoll.create(reader_id=debug.author_id, message_id = message.id)
                messagepoll.save()
                debug.reward=debug_value_form.cleaned_data['reward']
                debug.save()
                # 記錄系統事件
                if is_event_open(request) :                  
                    log = Log(user_id=request.user.id, event=u'評價除錯<'+ debug_value_form.cleaned_data['reward']+u'分><'+debug.author.first_name+'>')
                    log.save()                 
            except ObjectDoesNotExist:
			    pass
			    
        return redirect("/student/bug/"+str(bug_id))

			
def bug_detail(request, bug_id):
    bug = Bug.objects.get(id=bug_id)
    # List of active comments for this post
    debugs = Debug.objects.filter(bug_id=bug.id)
    datas = []
    for debug in debugs:
	    datas.append([debug, DebugValueForm(instance=debug)])
    try:
        filepath = request.FILES['file']
    except :
        filepath = False
    if request.method == 'POST':
        # A comment was posted
        debug_form = DebugForm(request.POST, request.FILES)
        if filepath :
            myfile = request.FILES['file']
            fs = FileSystemStorage()
            filename = "static/debug/"+str(request.user.id)+"/"+uuid4().hex
            fs.save(filename, myfile)				
        if debug_form.is_valid():
            # Create Comment object but don't save to database yet
            new_debug = debug_form.save(commit=False)
            # Assign the current bug to the comment
            new_debug.bug_id = bug.id
            new_debug.author_id = request.user.id
            new_debug.file = filename
            # Save the comment to the database
            new_debug.save()
            
            # create Message
            title = "<" + request.user.first_name + u">幫您除錯了<"+ bug.title + ">"
            url = request.get_full_path()
            message = Message.create(title=title, url=url, time=timezone.localtime(timezone.now()))
            message.save()
        
            messagepoll = MessagePoll.create(reader_id=bug.author_id, message_id = message.id)
            messagepoll.save()
            
            # 記錄系統事件
            bug = Bug.objects.get(id=bug_id)
            if is_event_open(request) :              
                log = Log(user_id=request.user.id, event=u'幫忙除錯<'+ bug.title +'>')
                log.save()               

            return redirect("/student/bug/"+bug_id)
    else:
        debug_form = DebugForm()      	
    return render_to_response('student/bug_detail.html',{'bug': bug,'datas': datas, 'debug_form': debug_form}, context_instance=RequestContext(request))
       
def bug_download(request, bug_id):
    bug = Bug.objects.get(id=bug_id)
    username = bug.author.first_name
    title = bug.title.replace(",", "")		
    filename = "Bug_"+username + "_" + title + ".sb2"
    download =  settings.BASE_DIR + "/" + str(bug.file)
    wrapper = FileWrapper(file( download, "r" ))
    response = HttpResponse(wrapper, content_type = 'application/force-download')
    #response = HttpResponse(content_type='application/force-download')
    response['Content-Disposition'] = 'attachment; filename={0}'.format(filename.encode('utf8'))
    # It's usually a good idea to set the 'Content-Length' header too.
    # You can also set any other required headers: Cache-Control, etc.
    return response

def debug_download(request, debug_id):
    debug = Debug.objects.get(id=debug_id)
    bug = Bug.objects.get(id=debug.bug_id)
    username = debug.author.first_name
    title = bug.title.replace(",", "")
    filename = "Debug_"+username + "_" + title + ".sb2"
    download =  settings.BASE_DIR + "/" + str(debug.file)
    wrapper = FileWrapper(file( download, "r" ))
    response = HttpResponse(wrapper, content_type = 'application/force-download')
    #response = HttpResponse(content_type='application/force-download')
    response['Content-Disposition'] = 'attachment; filename={0}'.format(filename.encode('utf8'))
    # It's usually a good idea to set the 'Content-Length' header too.
    # You can also set any other required headers: Cache-Control, etc.
    return response
	
class BugListClassView(ListView):
    context_object_name = 'bugs'
    paginate_by = 10
    template_name = 'student/bug_list.html'
    def get_queryset(self):
        class_bugs = []
        enrolls = Enroll.objects.filter(classroom_id=self.kwargs['classroom_id'])
        for enroll in enrolls:
            bugs = Bug.objects.filter(author_id=enroll.student_id, classroom_id=self.kwargs['classroom_id'])
            for bug in bugs:
                debugs = Debug.objects.filter(bug_id=bug.id)
                class_bugs.append([bug, debugs])
        def getKey(custom):
            return custom[0].publish	
        sorted_bug = sorted(class_bugs, key=getKey, reverse=True)
        # 記錄系統事件
        classroom_name = Classroom.objects.get(id=self.kwargs['classroom_id']).name
        if is_event_open(self.request) :          
            log = Log(user_id=self.request.user.id, event=u'查看所有除錯<'+ classroom_name +'>')
            log.save()         
        return sorted_bug
            
    
class BugCreateView(CreateView):
    model = Bug
    form_class = BugForm
    def form_valid(self, form):
        try:
            filepath = self.request.FILES['file']
        except :
            filepath = False
        if filepath :
                myfile = self.request.FILES['file']
                fs = FileSystemStorage()
                filename = "static/bug/"+str(self.request.user.id)+"/"+uuid4().hex
                fs.save(filename, myfile)
        self.object = form.save(commit=False)
        self.object.author_id = self.request.user.id
        self.object.classroom_id = self.kwargs['classroom_id']
        self.object.file = filename
        self.object.save()

        # create Message
        title = self.request.user.first_name + u"--提出一個Bug<" + form.cleaned_data['title'] + ">"
        url = "/student/bug/" + str(self.object.id)
        message = Message.create(title=title, url=url, time=timezone.now())
        message.save()
        
        enrolls = Enroll.objects.filter(student_id = self.request.user.id, classroom_id=self.kwargs['classroom_id'])
        for enroll in enrolls:
        # message for teacher
            messagepoll = MessagePoll.create(message_id = message.id,reader_id=enroll.classroom.teacher_id)
            messagepoll.save()
        # message for classsmates
            classmates = Enroll.objects.filter(classroom_id = enroll.classroom_id)
            for classmate in classmates:
				if not classmate.student_id == self.request.user.id:
					messagepoll = MessagePoll.create(message_id = message.id,reader_id=classmate.student_id)
					messagepoll.save()                
        # 記錄系統事件
        classroom_name = Classroom.objects.get(id=self.kwargs['classroom_id']).name        
        if is_event_open(self.request) :          
            log = Log(user_id=self.request.user.id, event=u'張貼Bug<'+ classroom_name +'>')
            log.save() 					
        return redirect(url)

# 說明作品編號
def work_help(request):
        return render_to_response('student/work_help.html', context_instance=RequestContext(request))



# 積分排行榜
class LoginLogListView(ListView):
    context_object_name = 'visitorlogs'
    paginate_by = 20
    template_name = 'student/login_log.html'
    def get_queryset(self):
        visitorlogs = VisitorLog.objects.filter(user_id=self.kwargs['user_id']).order_by("-id")
        # 記錄系統事件
        if is_event_open(self.request) :          
            user = User.objects.get(id=self.kwargs['user_id'])
            log = Log(user_id=self.request.user.id, event=u'查看登入記錄<'+user.first_name+'>')
            log.save()          
        return visitorlogs
        
    def get_context_data(self, **kwargs):
        context = super(LoginLogListView, self).get_context_data(**kwargs)
        if self.request.GET.get('page') :
            context['page'] = int(self.request.GET.get('page')) * 20 - 20
        else :
            context['page'] = 0
        return context        
        
# 積分排行榜
class NoteListView(ListView):
    context_object_name = 'notes'
    paginate_by = 20
    template_name = 'student/note.html'
    
    def get_queryset(self):
        notes = Note.objects.filter(user_id=self.request.user.id, classroom_id=0).order_by("-id")
        # 記錄系統事件
        if is_event_open(self.request) :          
            log = Log(user_id=self.request.user.id, event=u'查看學習筆記')
            log.save()          
        return notes
        
    def get_context_data(self, **kwargs):
        context = super(NoteListView, self).get_context_data(**kwargs)
        if self.request.GET.get('page') :
            context['page'] = int(self.request.GET.get('page')) * 20 - 20
        else :
            context['page'] = 0
        return context   
    
# 學習筆記匯出到word    
def doc_download(request):
    
    # 記錄系統事件
    if is_event_open(request) :       
        log = Log(user_id=request.user.id, event=u'下載學習筆記到Word')
        log.save()  

    document = Document()
    docx_title="Note-"+str(timezone.localtime(timezone.now()).date())+".docx"

    notes = Note.objects.filter(user_id=request.user.id, classroom_id=0).order_by("-id")
    paragraph = document.add_paragraph(request.user.first_name + u'的學習筆記')
    table = document.add_table(rows=1, cols=3)
    table.style = 'TableGrid'    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'課別'
    hdr_cells[1].text = u'日期'
    hdr_cells[2].text = u'內容'
    for note in notes:
        row_cells = table.add_row().cells
        row_cells[0].text = note.lesson
        row_cells[1].text = str(timezone.localtime(note.publication_date).strftime("%b %d %Y %H:%M:%S"))
        h = HTML2Text()
        h.ignore_links = True
        h.ignore_images = False

        text  = h.handle(note.memo)
        paragraph_number = 0
        while text.find("##[") != -1 :
            paragraph_number = paragraph_number + 1
            start = text.find("##[")
            row_cells[2].add_paragraph(text[:start])            
            end = text.find("]##")
            image_text = text[start:end+3]
            image = text[start+25:end]
            paragraph = row_cells[2].paragraphs[paragraph_number]
            run = paragraph.add_run()
            fh = open("imageToSave.png", "wb")
            fh.write(image.decode('base64'))
            fh.close()
            run.add_picture("imageToSave.png")
            text = string.replace(text, text[:start], "", 1)
            text = string.replace(text, image_text, "", 1)
        row_cells[2].add_paragraph(text)

    # Prepare document for download        
    # -----------------------------
    f = StringIO.StringIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length


    return response

# 影片
class VideoListView(ListView):
    context_object_name = 'videos'
    template_name = 'student/video.html'
    
    def get_queryset(self):
				videos = VideoLogHelper().getLogByUserid(self.kwargs['user_id'])
				# 記錄系統事件
				if is_event_open(self.request) :          
						log = Log(user_id=self.request.user.id, event=u'查看影片記錄')
						log.save()          
				return videos
        
    def get_context_data(self, **kwargs):
        context = super(VideoListView, self).get_context_data(**kwargs)
        context['video_url'] = video_url
        if self.request.GET.get('page') :
            context['page'] = int(self.request.GET.get('page')) * 20 - 20
        else :
            context['page'] = 0
        return context  

# 列出所有公告
class AnnounceListView(ListView):
    model = Message
    context_object_name = 'messages'
    template_name = 'student/announce_list.html'    
    paginate_by = 20
    
    def get_queryset(self):
        classroom = Classroom.objects.get(id=self.kwargs['classroom_id'])
        # 記錄系統事件
        if is_event_open(self.request) :    
            log = Log(user_id=self.request.user.id, event='查看班級公告')
            log.save()        
        messages = Message.objects.filter(classroom_id=classroom.id, author_id=classroom.teacher_id).order_by("-id")
        queryset = []
        for message in messages:
            try: 
                messagepoll = MessagePoll.objects.get(message_id=message.id, reader_id=self.request.user.id, classroom_id=classroom.id)
                queryset.append([messagepoll, message])
            except ObjectDoesNotExist :
                pass
        return queryset
        
    def get_context_data(self, **kwargs):
        context = super(AnnounceListView, self).get_context_data(**kwargs)
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        return context	    

    # 限本班同學
    def render_to_response(self, context):
        try:
            enroll = Enroll.objects.get(student_id=self.request.user.id, classroom_id=self.kwargs['classroom_id'])
        except ObjectDoesNotExist :
            return redirect('/')
        return super(AnnounceListView, self).render_to_response(context)        
			
# 列出分組12堂課所有作業
def work1(request, classroom_id):
        classroom_name = Classroom.objects.get(id=classroom_id).name
        lessons = []
        group = Enroll.objects.get(student_id=request.user.id, classroom_id=classroom_id).group
        for lesson in range(41):
          student_groups = []					         
          enrolls = Enroll.objects.filter(classroom_id=classroom_id, group=group)
          group_assistants = []
          assistants = []
          works = []
          scorer_name = ""
          for enroll in enrolls: 
              try:    
                  work = Work.objects.get(user_id=enroll.student_id, index=lesson+1)
                  if work.scorer > 0 :
                      scorer = User.objects.get(id=work.scorer)
                      scorer_name = scorer.first_name
                  else :
                      scorer_name = "X"
              except ObjectDoesNotExist:
                  work = Work(index=lesson, user_id=1)
              works.append([enroll, work.score, scorer_name, work.memo])
              try :
                  assistant = Assistant.objects.get(student_id=enroll.student.id, classroom_id=classroom_id, lesson=lesson+1)
                  group_assistants.append(enroll)
                  assistants.append(enroll.student_id)
              except ObjectDoesNotExist:
                  pass
          student_groups.append([group, works, group_assistants, assistants])
          lessons.append([lesson_list[lesson], student_groups])
        # 記錄系統事件
        if is_event_open(request) :            
            log = Log(user_id=request.user.id, event=u'查詢作業小老師<'+classroom_name+'>')
            log.save()         
        return render_to_response('student/work1.html', {'lessons':lessons, 'classroom_id':classroom_id}, context_instance=RequestContext(request))
						
# 日曆：班級登入列表
class LoginCalendarClassView(ListView):
    context_object_name = 'lists'
    #paginate_by = 50
    template_name = 'student/calendar.html'

    def get_queryset(self):    
        # 記錄系統事件
        classroom = Classroom.objects.get(id=self.kwargs['classroom_id'])
        enrolls = Enroll.objects.filter(classroom_id=classroom.id).order_by("seat")
        querysets = []
        log = Log(user_id=self.request.user.id, event=u'查看班級登入記錄<'+classroom.name+'>')
        log.save()
        for enroll in enrolls:
            user_logs = Log.objects.filter(user_id=enroll.student_id, event="登入系統").order_by("id")
            #weeklogs = groupby(user_logs, key=lambda row: (localtime(row.publish).isocalendar()[1]))
            logs = groupby(user_logs, key=lambda row: (localtime(row.publish).year, localtime(row.publish).month, localtime(row.publish).day))
            month_lists = []
            for key, value in logs:
                events = list(value)
                month_lists.append([key,events])
            if len(month_lists) > 0 :
                querysets.append([enroll, month_lists, (month_lists[-1][0][0]-month_lists[0][0][0])*150+200])
            else :
                querysets.append([enroll, month_lists,200])
        return querysets
        
    def get_context_data(self, **kwargs):
        context = super(LoginCalendarClassView, self).get_context_data(**kwargs)
        return context	
# 說明作品編號
def test(request):
        return render_to_response('student/test.html', context_instance=RequestContext(request))