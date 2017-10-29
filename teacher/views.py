# -*- coding: UTF-8 -*-
#from django.shortcuts import render
from django.shortcuts import render_to_response, redirect
from django.contrib.auth.models import User
from django.http import HttpResponse
#from django.contrib.auth import authenticate, login
from django.template import RequestContext
from django.contrib.auth.decorators import login_required
#from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.views.generic import ListView, DetailView, CreateView
from django.core.exceptions import ObjectDoesNotExist, MultipleObjectsReturned
#from django.contrib.auth.models import Group
from teacher.models import Classroom
from account.models import Log, Message, MessagePoll, Profile, Note, MessageFile
from student.models import Enroll, Work, EnrollGroup, Assistant, Exam, WorkFile
from show.models import Round
from .forms import ClassroomForm, ScoreForm,  CheckForm1, CheckForm2, CheckForm3, CheckForm4, AnnounceForm
#from django.views.generic.edit import ModelFormMixin
#from django.http import HttpResponseRedirect
import StringIO
from datetime import datetime
import xlsxwriter
from student.lesson import *
from account.avatar import *
from student.html2text import *
from account.models import Profile, PointHistory
from django.utils import timezone
from django.http import JsonResponse
from django.utils.timezone import localtime
from django.utils import timezone
from docx import *
from docx.shared import Inches
import string
import base64
import sys
import os
from django.conf import settings
from django.core.files.storage import FileSystemStorage
from uuid import uuid4
from wsgiref.util import FileWrapper#from django.contrib.auth.decorators import login_required, user_passes_test
from itertools import groupby
from account.helper import VideoLogHelper

# 判斷是否為授課教師
def is_teacher(user, classroom_id):
    return user.groups.filter(name='teacher').exists() and Classroom.objects.filter(teacher_id=user.id, id=classroom_id).exists()

# 判斷是否開啟事件記錄
def is_event_open(request):
        enrolls = Enroll.objects.filter(student_id=request.user.id)
        for enroll in enrolls:
            classroom = Classroom.objects.get(id=enroll.classroom_id)
            if classroom.event_open:
                return True
        return False

# 判斷是否開啟課程事件記錄
def is_event_video_open(request):
        enrolls = Enroll.objects.filter(student_id=request.user.id)
        for enroll in enrolls:
            classroom = Classroom.objects.get(id=enroll.classroom_id)
            if classroom.event_video_open:
                return True
        return False
        
# 列出所有課程
class ClassroomListView(ListView):
    model = Classroom
    context_object_name = 'classrooms'
    #paginate_by = 1000
    def get_queryset(self):
        # 記錄系統事件
        if is_event_open(self.request) :    
            log = Log(user_id=self.request.user.id, event='查看任課班級')
            log.save()        
        queryset = Classroom.objects.filter(teacher_id=self.request.user.id).order_by("-id")
        return queryset
        
#新增一個課程
class ClassroomCreateView(CreateView):
    model = Classroom
    form_class = ClassroomForm
    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.teacher_id = self.request.user.id
        self.object.event_open = False
        self.object.event_video_open = False
        self.object.save()
        # 將教師設為0號學生
        enroll = Enroll(classroom_id=self.object.id, student_id=self.request.user.id, seat=0)
        enroll.save()     
        # 新增1次班級創意秀
        show = Round(classroom_id=self.object.id)
        show.save()
        # 記錄系統事件
        if is_event_open(self.request) :            
            log = Log(user_id=self.request.user.id, event=u'新增任課班級<'+self.object.name+'>')
            log.save()                
        return redirect("/teacher/classroom")        
        
# 修改選課密碼
def classroom_edit(request, classroom_id):
    # 限本班任課教師
    if not is_teacher(request.user, classroom_id):
        return redirect("homepage")
    classroom = Classroom.objects.get(id=classroom_id)
    if request.method == 'POST':
        form = ClassroomForm(request.POST)
        if form.is_valid():
            classroom.name =form.cleaned_data['name']
            classroom.password = form.cleaned_data['password']
            classroom.save()
            # 記錄系統事件
            if is_event_open(request) :                
                log = Log(user_id=request.user.id, event=u'修改選課密碼<'+classroom.name+'>')
                log.save()                    
            return redirect('/teacher/classroom')
    else:
        form = ClassroomForm(instance=classroom)

    return render_to_response('form.html',{'form': form}, context_instance=RequestContext(request))        
    
# 退選
def unenroll(request, enroll_id, classroom_id):
    # 限本班任課教師
    if not is_teacher(request.user, classroom_id):
        return redirect("homepage")    
    enroll = Enroll.objects.get(id=enroll_id)
    enroll.delete()
    classroom_name = Classroom.objects.get(id=classroom_id).name
    # 記錄系統事件
    if is_event_open(request) :        
        log = Log(user_id=request.user.id, event=u'退選<'+classroom_name+'>')
        log.save()       
    return redirect('/student/classmate/'+classroom_id)  

# 列出班級所有作業
def work(request, classroom_id):
    # 限本班任課教師
    if not is_teacher(request.user, classroom_id):
        return redirect("homepage")    
    classroom = Classroom.objects.get(id=classroom_id)
    # 記錄系統事件
    if is_event_open(request) :    
        log = Log(user_id=request.user.id, event=u'列出班級所有作業<'+classroom.name+'>')
        log.save()              
    return render_to_response('teacher/work.html', {'lesson_list':lesson_list, 'classroom': classroom}, context_instance=RequestContext(request))

# 列出分組12堂課所有作業
def work1(request, classroom_id):
    # 限本班任課教師
    if not is_teacher(request.user, classroom_id):
        return redirect("homepage")    
    classroom_name = Classroom.objects.get(id=classroom_id).name
    lessons = []
    groups = [group for group in EnrollGroup.objects.filter(classroom_id=classroom_id)]
    enroll_pool = [enroll for enroll in Enroll.objects.filter(classroom_id=classroom_id).order_by('seat')]
    student_ids = map(lambda a: a.student_id, enroll_pool)
    work_pool = Work.objects.filter(user_id__in=student_ids)
    user_pool = [user for user in User.objects.filter(id__in=work_pool.values('scorer'))]
    assistant_pool = [assistant for assistant in Assistant.objects.filter(classroom_id=classroom_id)]
    for lesson in range(41):
        student_groups = []					
        for group in groups:
            members = filter(lambda u: u.group == group.id, enroll_pool)
            group_assistants = []
            works = []
            scorer_name = ""
            for member in members:
                work = filter(lambda w: w.index == lesson+1 and w.user_id == member.student_id, work_pool)
                if work:
                    work = work[0]
                    scorer = filter(lambda u: u.id == work.scorer, user_pool)
                    scorer_name = scorer[0].first_name if scorer else 'X'
                else:
                    work = Work(index=lesson+1, user_id=1)
                works.append([member, work.score, scorer_name, work.memo])
                assistant = filter(lambda a: a.student_id == member.student_id and a.lesson == lesson+1, assistant_pool)
                if assistant:
                    group_assistants.append(member)
            student_groups.append([group, works, group_assistants])
        lessons.append([lesson_list[lesson], student_groups])
    # 記錄系統事件
    if is_event_open(request) :            
        log = Log(user_id=request.user.id, event=u'以分組顯示作業<'+classroom_name+'>')
        log.save()         
    return render_to_response('teacher/work1.html', {'lessons':lessons, 'classroom_id':classroom_id}, context_instance=RequestContext(request))
			
			
# 列出某作業所有同學名單
def score(request, classroom_id, index):
    # 限本班任課教師
    if not is_teacher(request.user, classroom_id):
        return redirect("homepage")    
    enrolls = Enroll.objects.filter(classroom_id=classroom_id)
    classroom_name = Classroom.objects.get(id=classroom_id).name
    classmate_work = []
    scorer_name = ""
    for enroll in enrolls:
        try:    
            work = Work.objects.get(user_id=enroll.student_id, index=index)
            if work.scorer > 0 :
                scorer = User.objects.get(id=work.scorer)
                scorer_name = scorer.first_name
            else :
                scorer_name = "1"
        except ObjectDoesNotExist:
            work = Work(index=index, user_id=0)
        except MultipleObjectsReturned:
            work =  Work.objects.filter(user_id=enroll.student_id, index=index).order_by("-id")[0]
        try:
			group_name = EnrollGroup.objects.get(id=enroll.group).name
        except ObjectDoesNotExist:
			group_name = "沒有組別"
        assistant = Assistant.objects.filter(classroom_id=classroom_id, student_id=enroll.student_id, lesson=index)
        if assistant.exists():
            classmate_work.append([enroll,work,1, scorer_name, group_name])
        else :
            classmate_work.append([enroll,work,0, scorer_name, group_name])            
    lesson = lesson_list[int(index)-1]

    def getKey(custom):
        return custom[0].seat
	
    classmate_work = sorted(classmate_work, key=getKey)
    
    # 記錄系統事件
    if is_event_open(request) :        
        log = Log(user_id=request.user.id, event=u'列出某作業所有同學名單<'+classroom_name+'><'+index+'>')
        log.save()          
    return render_to_response('teacher/score.html',{'classmate_work': classmate_work, 'classroom_id':classroom_id, 'lesson':lesson, 'index': index}, context_instance=RequestContext(request))


# 教師評分
def scoring(request, classroom_id, user_id, index):
    user = User.objects.get(id=user_id)
    enroll = Enroll.objects.get(classroom_id=classroom_id, student_id=user_id)
    try:
        assistant = Assistant.objects.filter(classroom_id=classroom_id,lesson=index,student_id=request.user.id)
    except ObjectDoesNotExist:            
        if not is_teacher(request.user, classroom_id):
            return render_to_response('message.html', {'message':"您沒有權限"}, context_instance=RequestContext(request))
        
    try:
        work3 = Work.objects.get(user_id=user_id, index=index)
    except ObjectDoesNotExist:
        work3 = Work(index=index, user_id=user_id)
    except MultipleObjectsReturned:
        work3 =  Work.objects.filter(user_id=user_id, index=index).order_by("-id")[0]				
				
    workfiles = WorkFile.objects.filter(work_id=work3.id).order_by("-id")
		
    if request.method == 'POST':
        form = ScoreForm(request.user, request.POST)
        if form.is_valid():
            work = Work.objects.filter(index=index, user_id=user_id)
            if not work.exists():
                work = Work(index=index, user_id=user_id, score=form.cleaned_data['score'], publication_date=timezone.now())
                work.save()
                # 記錄系統事件
                if is_event_open() :            
                    log = Log(user_id=request.user.id, event=u'新增評分<'+user.first_name+'><'+work.score+'分>')
                    log.save()                      
            else:
                if work[0].score < 0 :   
                    # 小老師
                    if not is_teacher(request.user, classroom_id):
    	                # credit
                        update_avatar(request.user.id, 2, 1)
                        # History
                        history = PointHistory(user_id=request.user.id, kind=2, message='1分--小老師:<'+lesson_list[int(index)-1][2]+'><'+enroll.student.first_name.encode('utf-8')+'>', url=request.get_full_path())
                        history.save()				
    
				    # credit
                    update_avatar(enroll.student_id, 1, 1)
                    # History
                    history = PointHistory(user_id=user_id, kind=1, message='1分--作業受評<'+lesson_list[int(index)-1][2]+'><'+request.user.first_name.encode('utf-8')+'>', url=request.get_full_path())
                    history.save()		                        
                
                work.update(score=form.cleaned_data['score'])
                work.update(scorer=request.user.id)
                # 記錄系統事件
                if is_event_open(request) :                   
                    log = Log(user_id=request.user.id, event=u'更新評分<'+user.first_name+u'><'+str(work[0].score)+u'分>')
                    log.save()                    
						
            if is_teacher(request.user, classroom_id):         
                if form.cleaned_data['assistant']:
                    try :
					    assistant = Assistant.objects.get(student_id=user_id, classroom_id=classroom_id, lesson=index)
                    except ObjectDoesNotExist:
                        assistant = Assistant(student_id=user_id, classroom_id=classroom_id, lesson=index)
                        assistant.save()	
                        
                    # create Message
                    title = "<" + assistant.student.first_name.encode("utf-8") + u">擔任小老師<".encode("utf-8") + lesson_list[int(index)-1][2] + ">"
                    url = "/teacher/score_peer/" + str(index) + "/" + classroom_id + "/" + str(enroll.group) 
                    message = Message.create(title=title, url=url, time=timezone.now())
                    message.save()                        
                    
                    group = Enroll.objects.get(classroom_id=classroom_id, student_id=assistant.student_id).group
                    if group > 0 :
                        enrolls = Enroll.objects.filter(group = group)
                        for enroll in enrolls:
                            # message for group member
                            messagepoll = MessagePoll.create(message_id = message.id,reader_id=enroll.student_id)
                            messagepoll.save()
                    
                return redirect('/teacher/score/'+classroom_id+'/'+index)
            else: 
                return redirect('/teacher/score_peer/'+index+'/'+classroom_id+'/'+str(enroll.group))

    else:
        work = Work.objects.filter(index=index, user_id=user_id)
        if not work.exists():
            form = ScoreForm(user=request.user)
        else:
            form = ScoreForm(instance=work[0], user=request.user)
    lesson = lesson_list[int(index)-1]
    return render_to_response('teacher/scoring.html', {'form': form,'workfiles':workfiles, 'index': index, 'work':work3, 'student':user, 'classroom_id':classroom_id, 'lesson':lesson}, context_instance=RequestContext(request))

# 小老師評分名單
def score_peer(request, index, classroom_id, group):
    try:
        assistant = Assistant.objects.get(lesson=index, classroom_id=classroom_id, student_id=request.user.id)
    except ObjectDoesNotExist:
        return redirect("/student/group/work/"+index+"/"+classroom_id)

    enrolls = Enroll.objects.filter(classroom_id=classroom_id, group=group)
    lesson = ""
    classmate_work = []
    workfiles = []
    for enroll in enrolls:
        if not enroll.student_id == request.user.id : 
            scorer_name = ""
            try:    
                work = Work.objects.get(user_id=enroll.student.id, index=index)
                if work.scorer > 0 :
                    scorer = User.objects.get(id=work.scorer)
                    scorer_name = scorer.first_name
            except ObjectDoesNotExist:
                work = Work(index=index, user_id=1)
            workfiles = WorkFile.objects.filter(work_id=work.id)
            classmate_work.append([enroll.student,work,1, scorer_name])
        lesson = lesson_list[int(index)-1]
    # 記錄系統事件
    if is_event_open(request) :        
        log = Log(user_id=request.user.id, event=u'小老師評分名單<'+index+'><'+group+'>')
        log.save()    
    return render_to_response('teacher/score_peer.html',{'enrolls':enrolls, 'workfiles': workfiles, 'classmate_work': classmate_work, 'classroom_id':classroom_id, 'lesson':lesson, 'index': index}, context_instance=RequestContext(request))

# 設定為小老師
def assistant(request, classroom_id, user_id, lesson):
    # 限本班任課教師
    if not is_teacher(request.user, classroom_id):
        return redirect("homepage")    
    user = User.objects.get(id=user_id)
    assistant = Assistant(student_id=user_id, classroom_id=classroom_id, lesson=lesson)
    assistant.save()
    # 記錄系統事件
    if is_event_open(request) :        
        log = Log(user_id=request.user.id, event=u'設為小老師<'.encode("utf-8")+user.first_name.encode("utf-8")+'><'+ lesson_list[int(lesson)-1][2] + ">")
        log.save()    
    
    group = Enroll.objects.get(classroom_id=classroom_id, student_id=assistant.student_id).group
    # create Message
    title = "<" + assistant.student.first_name.encode("utf-8") + u">擔任小老師<".encode("utf-8") + lesson_list[int(lesson)-1][2] + ">"
    url = "/teacher/score_peer/" + str(lesson) + "/" + classroom_id + "/" + str(group) 
    message = Message.create(title=title, url=url, time=timezone.now())
    message.save()                        
        
    if group > 0 :
        enrolls = Enroll.objects.filter(group = group)
        for enroll in enrolls:
            # message for group member
            messagepoll = MessagePoll.create(message_id = message.id,reader_id=enroll.student_id)
            messagepoll.save()
    
    return redirect('/teacher/score/'+str(assistant.classroom_id)+"/"+lesson)    
    
# 取消小老師
def assistant_cancle(request, classroom_id, user_id, lesson):
    # 限本班任課教師
    if not is_teacher(request.user, classroom_id):
        return redirect("homepage")    
    user = User.objects.get(id=user_id)   
    assistant = Assistant.objects.get(student_id=user_id, classroom_id=classroom_id, lesson=lesson)
    assistant.delete()
    # 記錄系統事件
    if is_event_open(request) :        
        log = Log(user_id=request.user.id, event=u'取消小老師<'.encode("utf-8")+user.first_name.encode("utf-8")+'><'+ lesson_list[int(lesson)-1][2] + ">")
        log.save()     
    
    # create Message
    title = "<" + assistant.student.first_name.encode("utf-8") + u">取消小老師<".encode("utf-8") + lesson_list[int(lesson)-1][2] + ">"
    url = "/student/group/work/" + str(lesson) + "/" + classroom_id 
    message = Message.create(title=title, url=url, time=timezone.now())
    message.save()                        
        
    group = Enroll.objects.get(classroom_id=classroom_id, student_id=assistant.student_id).group
    if group > 0 :
        enrolls = Enroll.objects.filter(group = group)
        for enroll in enrolls:
            # message for group member
            messagepoll = MessagePoll.create(message_id = message.id,reader_id=enroll.student_id)
            messagepoll.save()
    

    return redirect('/teacher/score/'+str(assistant.classroom_id)+"/"+lesson)    
    
# 以分組顯示作業
def work_group(request, lesson, classroom_id):
        # 限本班任課教師
        if not is_teacher(request.user, classroom_id):
            return redirect("homepage")    
        classroom_name = Classroom.objects.get(id=classroom_id).name
        student_groups = []
        groups = EnrollGroup.objects.filter(classroom_id=classroom_id)
        try:
                student_group = Enroll.objects.get(student_id=request.user.id, classroom_id=classroom_id).group
        except ObjectDoesNotExist :
                student_group = []		
        for group in groups:
            enrolls = Enroll.objects.filter(classroom_id=classroom_id, group=group.id)
            group_assistants = []
            works = []
            scorer_name = ""
            for enroll in enrolls: 
                try:    
                    work = Work.objects.get(user_id=enroll.student_id, index=lesson)
                    if work.scorer > 0 :
                        scorer = User.objects.get(id=work.scorer)
                        scorer_name = scorer.first_name
                    else :
                        scorer_name = "X"
                except ObjectDoesNotExist:
                    work = Work(index=lesson, user_id=1)
                works.append([enroll, work.score, scorer_name])
                try :
                    assistant = Assistant.objects.get(student_id=enroll.student.id, classroom_id=classroom_id, lesson=lesson)
                    group_assistants.append(enroll)
                except ObjectDoesNotExist:
				    pass
            student_groups.append([group, works, group_assistants])
        lesson_data = lesson_list[int(lesson)-1]		
        # 記錄系統事件
        if is_event_open(request) :            
            log = Log(user_id=request.user.id, event=u'以分組顯示作業<'+lesson+'><'+classroom_name+'>')
            log.save()         
        return render_to_response('teacher/work_group.html', {'lesson':lesson, 'lesson_data':lesson_data, 'student_groups':student_groups, 'classroom_id':classroom_id, 'student_group':student_group}, context_instance=RequestContext(request))

# 心得
def memo(request, classroom_id):
        # 限本班任課教師
        if not is_teacher(request.user, classroom_id):
            return redirect("homepage")    
        enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by("seat")
        classroom_name = Classroom.objects.get(id=classroom_id).name
        # 記錄系統事件
        if is_event_open(request) :            
            log = Log(user_id=request.user.id, event=u'查閱心得<'+classroom_name+'>')
            log.save()  
        return render_to_response('teacher/memo.html', {'enrolls':enrolls, 'classroom_name':classroom_name}, context_instance=RequestContext(request))

# 評分某同學某進度心得
@login_required
def check(request, user_id, unit,classroom_id):
    # 限本班任課教師
    if not is_teacher(request.user, classroom_id):
        return redirect("homepage")

    user_name = User.objects.get(id=user_id).first_name
    del lesson_list[:]
    reset()
    works = Work.objects.filter(user_id=user_id)
    for work in works:
        lesson_list[work.index-1].append(work.score)
        lesson_list[work.index-1].append(work.publication_date)
        if work.score > 0 :
            score_name = User.objects.get(id=work.scorer).first_name
            lesson_list[work.index-1].append(score_name)
        else :
            lesson_list[work.index-1].append("尚未評分!")
        lesson_list[work.index-1].append(work.memo)
    c = 0
    for lesson in lesson_list:
        assistant = Assistant.objects.filter(student_id=user_id, lesson=c+1)
        if assistant.exists() :
            lesson.append("V")
        else :
            lesson.append("")
        c = c + 1
    user = User.objects.get(id=user_id)

    if unit == "1" :
        if request.method == 'POST':
            form = CheckForm1(request.POST)
            if form.is_valid():
                enroll = Enroll.objects.get(student_id=user_id, classroom_id=classroom_id)
                enroll.score_memo1=form.cleaned_data['score_memo1']
                enroll.save()
                
                # 記錄系統事件
                if is_event_open(request) :                    
                    log = Log(user_id=request.user.id, event=u'批改12堂課心得<'+user_name+'>')
                    log.save()                  
						
                if form.cleaned_data['certificate']:		
                    return redirect('/certificate/make_certification/'+unit+'/'+str(enroll.id)+'/certificate')
                else:
                    return redirect('/teacher/memo/'+classroom_id)
        else:
            enroll = Enroll.objects.get(student_id=user_id, classroom_id=classroom_id)
            form = CheckForm1(instance=enroll)
    elif unit == "2":
        if request.method == 'POST':
            form = CheckForm2(request.POST)
            if form.is_valid():
                enroll = Enroll.objects.get(student_id=user_id, classroom_id=classroom_id)
                enroll.score_memo2=form.cleaned_data['score_memo2']
                enroll.save()

                # 記錄系統事件
                if is_event_open(request) :                    
                    log = Log(user_id=request.user.id, event=u'批改實戰入門心得<'+user_name+'>')
                    log.save()      
						
                if form.cleaned_data['certificate']:		
                    return redirect('/certificate/make_certification/'+unit+'/'+str(enroll.id)+'/certificate')
                else:
                    return redirect('/teacher/memo/'+classroom_id)					
        else:
            enroll = Enroll.objects.get(student_id=user_id, classroom_id=classroom_id)
            form = CheckForm2(instance=enroll)
    elif unit == "3":
        if request.method == 'POST':
            form = CheckForm3(request.POST)
            if form.is_valid():
                enroll = Enroll.objects.get(student_id=user_id, classroom_id=classroom_id)
                enroll.score_memo3=form.cleaned_data['score_memo3']
                enroll.save()
                # 記錄系統事件
                if is_event_open(request) :                    
                    log = Log(user_id=request.user.id, event=u'批改實戰進擊心得<'+user_name+'>')
                    log.save() 
						
                if form.cleaned_data['certificate']:		
                    return redirect('/certificate/make_certification/'+unit+'/'+str(enroll.id)+'/certificate')
                else:
                    return redirect('/teacher/memo/'+classroom_id)							
        else:
            enroll = Enroll.objects.get(student_id=user_id, classroom_id=classroom_id)
            form = CheckForm3(instance=enroll)

    else:
        if request.method == 'POST':
            form = CheckForm4(request.POST)
            if form.is_valid():
                enroll = Enroll.objects.get(student_id=user_id, classroom_id=classroom_id)
                enroll.score_memo4=form.cleaned_data['score_memo4']
                enroll.save()
                # 記錄系統事件
                if is_event_open(request) :                    
                    log = Log(user_id=request.user.id, event=u'批改實戰高手心得<'+user_name+'>')
                    log.save() 						
                if form.cleaned_data['certificate']:		
                    return redirect('/certificate/make_certification/'+unit+'/'+str(enroll.id)+'/certificate')
                else:
                    return redirect('/teacher/memo/'+classroom_id)							
        else:
            enroll = Enroll.objects.get(student_id=user_id, classroom_id=classroom_id)
            form = CheckForm4(instance=enroll)	
    # 記錄系統事件
    if is_event_open(request) :        
        log = Log(user_id=request.user.id, event=u'查閱個人心得<'+user_name+'>')
        log.save()  
    return render_to_response('teacher/check.html', {'form':form, 'works':works, 'lesson_list':lesson_list, 'student': user, 'unit':unit, 'classroom_id':classroom_id}, context_instance=RequestContext(request))

	
# 查閱全班測驗卷成績
def exam_list(request, classroom_id):
        # 限本班任課教師
        if not is_teacher(request.user, classroom_id):
            return redirect("homepage")
        enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by("seat")
        classroom_name=""
        enroll_exam = []		
        for enroll in enrolls:
            classroom_name = enroll.classroom.name
            exam_list = []
            for exam_id in range(3):
                exams = Exam.objects.filter(student_id=enroll.student_id, exam_id=exam_id+1)
                total = 0
                times = 0
                for exam in exams:
                    total += exam.score
                    times += 1
                exam_list.append(total)
                exam_list.append(times)
            enroll_exam.append([enroll, exam_list])
        # 記錄系統事件
        if is_event_open(request) :            
            log = Log(user_id=request.user.id, event=u'查閱測驗卷成績<'+classroom_name+'>')
            log.save() 	
        return render_to_response('teacher/exam_list.html', {'classroom_id':classroom_id, 'classroom_name':classroom_name, 'enroll_exam':enroll_exam}, context_instance=RequestContext(request))

# 查詢某項測驗的所有資料
def exam_detail(request, classroom_id, student_id, exam_id):
        # 限本班任課教師
        if not is_teacher(request.user, classroom_id):
            return redirect("homepage")
        exams = Exam.objects.filter(student_id=student_id, exam_id=exam_id)
        enroll = Enroll.objects.get(classroom_id=classroom_id, student_id=student_id)
        # 記錄系統事件
        if is_event_open(request) :            
            log = Log(user_id=request.user.id, event=u'檢視測驗資料<'+exam_id+'><'+enroll.student.first_name+'>')
            log.save() 	        
        return render_to_response('teacher/exam_detail.html', {'exams': exams, 'enroll':enroll}, context_instance=RequestContext(request))
		
# 結算成績
@login_required
def grade(request, classroom_id):
        # 限本班任課教師
        if not is_teacher(request.user, classroom_id):
            return redirect("homepage")
        classroom = Classroom.objects.get(id=classroom_id)
        # 記錄系統事件
        if is_event_open(request) :    
            log = Log(user_id=request.user.id, event=u'結算成績<'+classroom.name+'>')
            log.save() 	                
        return render_to_response('teacher/grade.html', {'classroom': classroom}, context_instance=RequestContext(request))

@login_required
def grade_unit1(request, classroom_id):
        # 限本班任課教師
        if not is_teacher(request.user, classroom_id):
            return redirect("homepage")    
        classroom = Classroom.objects.get(id=classroom_id)
        if not request.user.id == classroom.teacher_id:
            return redirect("/")
        enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by('seat')
        enroll_group = {}
        data = []
        for enroll in enrolls:
            enroll_score = []
            enroll_grade = []
            # 17個作品
            for i in range(17):
                try :
                    work = Work.objects.get(user_id=enroll.student_id, index=i+1)
                    if work.score == -1 :
                        score=0
                    else :
                        score = work.score
                    enroll_score.append(work.score)
                    enroll_grade.append(score/25.0)
                except ObjectDoesNotExist:
                    enroll_score.append('缺')
                    enroll_grade.append(0)
            # 3個測驗
            for i in range(3):
                exams = Exam.objects.filter(student_id=enroll.student_id, exam_id=i+1)
                if not exams.exists():
                    enroll_score.append(0)
                    enroll_grade.append(0)					
                else:
                    total = 0
                    times = 0
                    for exam in exams:
                        total += exam.score
                        times += 1
                    if times > 0 :
                        average = total/times
                        enroll_score.append(average)
                        enroll_grade.append(average/25.0)					
                    else:
                        enroll_score.append(0)
                        enroll_grade.append(0)
            # 1個心得
            enroll_score.append(enroll.score_memo1)
            enroll_grade.append(enroll.score_memo1/10.0)
			      # 1個積分
            profile = Profile.objects.get(user=enroll.student)
            point = profile.work + profile.assistant + profile.creative + profile.debug
            enroll_score.append(point)
            if point * 0.12 > 10:
                enroll_grade.append(10)
            else :
                enroll_grade.append(point*0.12)
            score = int(sum(enroll_grade))
				
            
            if enroll_group.has_key(enroll.group):
                enroll_group[enroll.group].append([enroll.student_id, score])
            else:
                enroll_group[enroll.group] = [[enroll.student_id, score]]        

            data.append([enroll, enroll_score, enroll_grade, score])            
     
        i = 0
        for enroll in enrolls:
            group_total = 0
            for group in enroll_group[enroll.group]:
                group_total += group[1]
            group_people = len(enroll_group[enroll.group])
            group_score = group_total / group_people
            data[i].append(group_score)
            i = i + 1 

        # 記錄系統事件
        if is_event_open(request) :    
            log = Log(user_id=request.user.id, event=u'查看成績<12堂課><'+classroom.name+'>')
            log.save() 	
        return render_to_response('teacher/grade_unit1.html', {'lesson_list':lesson_list, 'classroom':classroom, 'data':data}, context_instance=RequestContext(request))

@login_required
def grade_unit2(request, classroom_id):
        # 限本班任課教師
        if not is_teacher(request.user, classroom_id):
            return redirect("homepage")  
        enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by('seat')
        classroom = Classroom.objects.get(id=classroom_id)        
        data = []
        enroll_group = {}        
        for enroll in enrolls:
            enroll_score = []
            enroll_grade = []
            # 8個作品
            for i in range(8):
                try :
                    work = Work.objects.get(user_id=enroll.student_id, index=i+1+17)
                    if work.score == -1 :
                        score=0
                    else :
                        score = work.score
                    enroll_score.append(work.score)
                    enroll_grade.append(score/10)
                except ObjectDoesNotExist:
                    enroll_score.append('缺')
                    enroll_grade.append(0)
            # 1個心得
            enroll_score.append(enroll.score_memo2)
            enroll_grade.append(enroll.score_memo2/10.0)
			# 1個積分
            profile = Profile.objects.get(user=enroll.student)
            point = profile.work + profile.assistant + profile.creative + profile.debug
            enroll_score.append(point)
            if point * 0.1 > 10:
                enroll_grade.append(10)
            else :
                enroll_grade.append(point*0.1)
            score = int(sum(enroll_grade))
				
            if enroll_group.has_key(enroll.group):
                enroll_group[enroll.group].append([enroll.student_id, score])
            else:
                enroll_group[enroll.group] = [[enroll.student_id, score]]        

            data.append([enroll, enroll_score, enroll_grade, score])            
     
        i = 0
        for enroll in enrolls:
            group_total = 0
            for group in enroll_group[enroll.group]:
                group_total += group[1]
            group_people = len(enroll_group[enroll.group])
            group_score = group_total / group_people
            data[i].append(group_score)
            i = i + 1
            
        # 記錄系統事件
        if is_event_open(request) :    
            log = Log(user_id=request.user.id, event=u'查看成績<實戰入門><'+classroom.name+'>')
            log.save() 

        return render_to_response('teacher/grade_unit2.html', {'lesson_list':lesson_list, 'classroom':classroom, 'data':data}, context_instance=RequestContext(request))

@login_required
def grade_unit3(request, classroom_id):
        # 限本班任課教師
        if not is_teacher(request.user, classroom_id):
            return redirect("homepage")     
        enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by('seat')
        classroom = Classroom.objects.get(id=classroom_id)        
        data = []
        enroll_group = {}        
        for enroll in enrolls:
            enroll_score = []
            enroll_grade = []
            # 8個作品
            for i in range(8):
                try :
                    work = Work.objects.get(user_id=enroll.student_id, index=i+1+17+8)
                    if work.score == -1 :
                        score=0
                    else :
                        score = work.score
                    enroll_score.append(work.score)
                    enroll_grade.append(score/10)
                except ObjectDoesNotExist:
                    enroll_score.append('缺')
                    enroll_grade.append(0)
            # 1個心得
            enroll_score.append(enroll.score_memo3)
            enroll_grade.append(enroll.score_memo3/10.0)
			# 1個積分
            profile = Profile.objects.get(user=enroll.student)
            point = profile.work + profile.assistant + profile.creative + profile.debug
            enroll_score.append(point)
            if point * 0.08 > 10:
                enroll_grade.append(10)
            else :
                enroll_grade.append(point*0.08)
            score = int(sum(enroll_grade))
				
            if enroll_group.has_key(enroll.group):
                enroll_group[enroll.group].append([enroll.student_id, score])
            else:
                enroll_group[enroll.group] = [[enroll.student_id, score]]        

            data.append([enroll, enroll_score, enroll_grade, score])            
     
        i = 0
        for enroll in enrolls:
            group_total = 0
            for group in enroll_group[enroll.group]:
                group_total += group[1]
            group_people = len(enroll_group[enroll.group])
            group_score = group_total / group_people
            data[i].append(group_score)
            i = i + 1 

        # 記錄系統事件
        if is_event_open(request) :            
            log = Log(user_id=request.user.id, event=u'查看成績<實戰進擊><'+classroom.name+'>')
            log.save() 

        return render_to_response('teacher/grade_unit3.html', {'lesson_list':lesson_list, 'classroom':classroom, 'data':data}, context_instance=RequestContext(request))

@login_required
def grade_unit4(request, classroom_id):
        # 限本班任課教師
        if not is_teacher(request.user, classroom_id):
            return redirect("homepage")     
        enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by('seat')
        classroom = Classroom.objects.get(id=classroom_id)        
        data = []
        enroll_group = {}
        for enroll in enrolls:
            enroll_score = []
            enroll_grade = []
            # 8個作品
            for i in range(8):
                try :
                    work = Work.objects.get(user_id=enroll.student_id, index=i+1+17+8+8)
                    if work.score == -1 :
                        score=0
                    else :
                        score = work.score
                    enroll_score.append(work.score)
                    enroll_grade.append(score/10)
                except ObjectDoesNotExist:
                    enroll_score.append("缺")
                    enroll_grade.append(0)
            # 1個心得
            enroll_score.append(enroll.score_memo4)
            enroll_grade.append(enroll.score_memo4/10.0)
			      # 1個積分
            profile = Profile.objects.get(user=enroll.student)
            point = profile.work + profile.assistant + profile.creative + profile.debug
            enroll_score.append(point)
            if point * 0.06 > 10:
                enroll_grade.append(10)
            else :
                enroll_grade.append(point*0.06)
            score = int(sum(enroll_grade))
            
            if enroll_group.has_key(enroll.group):
                enroll_group[enroll.group].append([enroll.student_id, score])
            else:
                enroll_group[enroll.group] = [[enroll.student_id, score]]        

            data.append([enroll, enroll_score, enroll_grade, score])            
     
        i = 0
        for enroll in enrolls:
            group_total = 0
            for group in enroll_group[enroll.group]:
                group_total += group[1]
            group_people = len(enroll_group[enroll.group])
            group_score = group_total / group_people
            data[i].append(group_score)
            i = i + 1 
            
        # 記錄系統事件
        if is_event_open(request) :            
            log = Log(user_id=request.user.id, event=u'查看成<績實戰高手><'+classroom.name+'>')
            log.save() 
            
        return render_to_response('teacher/grade_unit4.html', {'enroll_group':enroll_group, 'lesson_list':lesson_list, 'classroom':classroom, 'data':data}, context_instance=RequestContext(request))

# 列出所有公告
class AnnounceListView(ListView):
    model = Message
    context_object_name = 'messages'
    template_name = 'teacher/announce_list.html'    
    paginate_by = 20
    def get_queryset(self):

        # 記錄系統事件
        if is_event_open(self.request) :    
            log = Log(user_id=self.request.user.id, event='查看班級公告')
            log.save()        
        queryset = Message.objects.filter(classroom_id=self.kwargs['classroom_id'], author_id=self.request.user.id).order_by("-id")
        return queryset
        
    def get_context_data(self, **kwargs):
        context = super(AnnounceListView, self).get_context_data(**kwargs)
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        return context	    

    # 限本班任課教師        
    def render_to_response(self, context):
        if not is_teacher(self.request.user, self.kwargs['classroom_id']):
            return redirect('/')
        return super(AnnounceListView, self).render_to_response(context)        
        
#新增一個公告
class AnnounceCreateView(CreateView):
    model = Message
    form_class = AnnounceForm
    template_name = 'teacher/announce_form.html'     
    def form_valid(self, form):
        classrooms = self.request.POST.getlist('classrooms')
        files = self.request.FILES.getlist('files')
        self.object = form.save(commit=False)
        filenames = []
        for file in files:
            fs = FileSystemStorage()
            filename = uuid4().hex
            fs.save("static/message/"+filename, file)
            filenames.append([filename, file.name])
        for classroom_id in classrooms:
            message = Message()
            message.title = u"[公告]" + self.object.title
            message.author_id = self.request.user.id	
            message.classroom_id = classroom_id
            message.content = self.object.content
            message.save()
            message.url = "/teacher/announce/detail/" + str(message.id)
            message.save()
            for filename in filenames:
                    messagefile = MessageFile(message_id=message.id, filename=filename[0], before_name=filename[1])
                    messagefile.save()
            # 班級學生訊息
            enrolls = Enroll.objects.filter(classroom_id=classroom_id)
            for enroll in enrolls:
                messagepoll = MessagePoll(message_id=message.id, reader_id=enroll.student_id, classroom_id=classroom_id)
                messagepoll.save()
        # 記錄系統事件
        if is_event_open(self.request) :            
            log = Log(user_id=self.request.user.id, event=u'新增公告<'+self.object.title+'>')
            log.save()                
        return redirect("/teacher/announce/"+self.kwargs['classroom_id'])       
        
    def get_context_data(self, **kwargs):
        context = super(AnnounceCreateView, self).get_context_data(**kwargs)
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['classrooms'] = Classroom.objects.filter(teacher_id=self.request.user.id).order_by("-id")
        return context	   
        
    # 限本班任課教師        
    def render_to_response(self, context):
        if not is_teacher(self.request.user, self.kwargs['classroom_id']):
            return redirect('/')
        return super(AnnounceCreateView, self).render_to_response(context)          
        
# 查看藝郎某項目
def announce_detail(request, message_id):
    message = Message.objects.get(id=message_id)
    classroom = Classroom.objects.get(id=message.classroom_id)
    
    announce_reads = []
    
    messagepolls = MessagePoll.objects.filter(message_id=message_id, classroom_id=classroom.id)
    for messagepoll in messagepolls:
        enroll = Enroll.objects.get(classroom_id=message.classroom_id, student_id=messagepoll.reader_id)
        announce_reads.append([enroll.seat, enroll.student.first_name, messagepoll])
    
    def getKey(custom):
        return custom[0]	
    announce_reads = sorted(announce_reads, key=getKey)
    
    if is_event_open(request) :            
        log = Log(user_id=request.user.id, event=u'查看公告<'+message.title+'>')
        log.save()

    files = MessageFile.objects.filter(message_id=message_id)
    return render_to_response('teacher/announce_detail.html', {'files':files, 'message':message, 'classroom':classroom, 'announce_reads':announce_reads}, context_instance=RequestContext(request))

def announce_download(request, messagefile_id):
    messagefile = MessageFile.objects.get(id=messagefile_id)
    download =  settings.BASE_DIR + "/static/message/" + messagefile.filename
    wrapper = FileWrapper(file( download, "r" ))
    filename = messagefile.before_name
    response = HttpResponse(wrapper, content_type = 'application/force-download')
    #response = HttpResponse(content_type='application/force-download')
    response['Content-Disposition'] = 'attachment; filename={0}'.format(filename.encode('utf8'))
    # It's usually a good idea to set the 'Content-Length' header too.
    # You can also set any other required headers: Cache-Control, etc.
    return response
    #return render_to_response('student/download.html', {'download':download})

# 記錄系統事件
class EventListView(ListView):
    context_object_name = 'events'
    paginate_by = 50
    template_name = 'teacher/event_list.html'

    def get_queryset(self):    
        classroom = Classroom.objects.get(id=self.kwargs['classroom_id'])
        # 記錄系統事件
        if is_event_open(self.request) :           
            log = Log(user_id=self.request.user.id, event=u'查看班級事件<'+classroom.name+'>')
            log.save()       
        enrolls = Enroll.objects.filter(classroom_id=self.kwargs['classroom_id']);
        users = []
        for enroll in enrolls:
            if enroll.seat > 0 :
                users.append(enroll.student_id)
        if self.kwargs['user_id'] == "0":
            if self.request.GET.get('q') != None:
                queryset = Log.objects.filter(user_id__in=users, event__icontains=self.request.GET.get('q')).order_by('-id')
            else :
                queryset = Log.objects.filter(user_id__in=users).order_by('-id')
        else :
            if self.request.GET.get('q') != None:
                queryset = Log.objects.filter(user_id=self.kwargs['user_id'],event__icontains=self.request.GET.get('q')).order_by('-id')
            else : 
                queryset = Log.objects.filter(user_id__in=users).order_by('-id')
        return queryset
        
    def get_context_data(self, **kwargs):
        context = super(EventListView, self).get_context_data(**kwargs)
        q = self.request.GET.get('q')
        context.update({'q': q})
        classroom = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['classroom'] = classroom
        context['is_event_open'] = Classroom.objects.get(id=self.kwargs['classroom_id']).event_open
        context['is_event_video_open'] = Classroom.objects.get(id=self.kwargs['classroom_id']).event_video_open       
        return context		
		
# 記錄系統事件
class Event12ListView(ListView):
    context_object_name = 'events'
    template_name = 'teacher/event12_list.html'

    def get_queryset(self):    
        classroom = Classroom.objects.get(id=self.kwargs['classroom_id'])
        # 記錄系統事件
        if is_event_open(self.request) :           
            log = Log(user_id=self.request.user.id, event=u'查看班級12堂課事件<'+classroom.name+'>')
            log.save()       
        enrolls = [enroll for enroll in Enroll.objects.filter(classroom_id=self.kwargs['classroom_id'], seat__gt=0).order_by("seat")]
        student_ids = [enroll.student_id for enroll in enrolls] # 取得所有修課學生的 id
        # 一次取得所有修課學生查看課程內容的紀錄
        logs = [log for log in Log.objects.filter(user_id__in=student_ids, event__regex='查看課程內容<\d+>').order_by('-id')]
        regex = re.compile(u'查看課程內容<(\d+)>')        
        events = []
        student_log = {}
        # 初始化 student_log
        for enroll in enrolls:
            student_log[enroll.student_id] = {}
            for i in range(12):
                student_log[enroll.student_id][i+1] = [] 
        # 將紀錄填入相對應的 student_log 位置
        for log in logs:
            cid = int(regex.match(log.event).group(1)) # 取出課程編號 1~11
            student_log[log.user_id][cid].append(log)
        # 計算每位學生每個課程的耗用時間
        for enroll in enrolls:
            durations = []
            for key in student_log[enroll.student_id]:
                log = student_log[enroll.student_id][key]
                if len(log) > 0:
                    duration = str(log[0].publish - log[-1].publish)
                    durations.append(duration.rsplit(':', 1)[0])
                else:
                    durations.append(0)
            events.append({'seat': enroll.seat, 'sid': enroll.student_id, 'durations': durations})
        return events
        
    def get_context_data(self, **kwargs):
        context = super(Event12ListView, self).get_context_data(**kwargs)
        classroom = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['classroom'] = classroom
        return context	

# 記錄系統事件
class Event12LessonListView(ListView):
    context_object_name = 'events'
    template_name = 'teacher/event12Lesson_list.html'

    def get_queryset(self):    
        # 記錄系統事件
        if is_event_open(self.request) :           
            log = Log(user_id=self.request.user.id, event=u'查看課程事件<'+self.kwargs['user_id']+'><'+self.kwargs['lesson']+'>')
            log.save()       
        queryset = Log.objects.filter(user_id=self.kwargs['user_id'],event__icontains="查看課程內容<"+self.kwargs['lesson']+">").order_by('-id')
        return queryset
        
    def get_context_data(self, **kwargs):
        context = super(Event12LessonListView, self).get_context_data(**kwargs)
        context['user_id'] = self.kwargs['user_id']
        context['lesson'] = self.kwargs['lesson']		
        return context			
	
def clear(request, classroom_id):
    Log.objects.all().delete()
    # 記錄系統事件
    if is_event_open(request) :       
        log = Log(user_id=request.user.id, event=u'清除所有事件')
        log.save()            
    return redirect("/account/event/0")
    
def event_excel(request, classroom_id):
    classroom = Classroom.objects.get(id=classroom_id)
    # 記錄系統事件
    if is_event_open(request) :       
        log = Log(user_id=request.user.id, event=u'下載事件到Excel')
        log.save()        
    output = StringIO.StringIO()
    workbook = xlsxwriter.Workbook(output)    
    #workbook = xlsxwriter.Workbook('hello.xlsx')
    worksheet = workbook.add_worksheet()
    date_format = workbook.add_format({'num_format': 'dd/mm/yy hh:mm:ss'})
    enrolls = Enroll.objects.filter(classroom_id=classroom_id);
    users = []
    for enroll in enrolls:
        if enroll.seat > 0 :
            users.append(enroll.student_id)
    events = Log.objects.filter(user_id__in=users).order_by('-id')
    index = 1
    for event in events:
        if event.user_id > 0 :
            worksheet.write('A'+str(index), event.user.first_name)
        else: 
            worksheet.write('A'+str(index), u'匿名')
        worksheet.write('B'+str(index), event.event)
        worksheet.write('C'+str(index), str(localtime(event.publish)))
        index = index + 1

    workbook.close()
    # xlsx_data contains the Excel file
    response = HttpResponse(content_type='application/vnd.ms-excel')
    response['Content-Disposition'] = 'attachment; filename=Report-'+classroom.name+'-'+str(localtime(timezone.now()).date())+'.xlsx'
    xlsx_data = output.getvalue()
    response.write(xlsx_data)
    return response

def event_make(request):
    action = request.POST.get('action')
    classroom_id = request.POST.get('classroomid')    
    if classroom_id and action :
            classroom = Classroom.objects.get(id=classroom_id)
            if action == 'open':
                classroom.event_open = True
            else :
                classroom.event_open = False
            classroom.save()
            return JsonResponse({'status':'ok'}, safe=False)
    else:
            return JsonResponse({'status':'ko'}, safe=False)
     
def event_video_make(request):
    action = request.POST.get('action')
    classroom_id = request.POST.get('classroomid')      
    if classroom_id and action :
            classroom = Classroom.objects.get(id=classroom_id)
            if action == 'open':
                classroom.event_video_open = True
            else :
                classroom.event_video_open = False
            classroom.save()
            return JsonResponse({'status':'ok'}, safe=False)
    else:
            return JsonResponse({'status':'ko'}, safe=False)

		
# 列出所有公告
class NoteListView(ListView):
    model = Note
    context_object_name = 'notes'
    template_name = 'teacher/note.html'    
    paginate_by = 20
    def get_queryset(self):

        # 記錄系統事件
        if is_event_open(self.request) :    
            log = Log(user_id=self.request.user.id, event='查看班級筆記')
            log.save()        
        queryset = Note.objects.filter(classroom_id=self.kwargs['classroom_id'], user_id=self.request.user.id).order_by("-id")
        return queryset
        
    def get_context_data(self, **kwargs):
        context = super(NoteListView, self).get_context_data(**kwargs)
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        return context	    

    # 限本班任課教師        
    def render_to_response(self, context):
        if not is_teacher(self.request.user, self.kwargs['classroom_id']):
            return redirect('/')
        return super(NoteListView, self).render_to_response(context)        
        
# 教學筆記匯出到word    
def doc_download(request, classroom_id):
    classroom = Classroom.objects.get(id=classroom_id)
    
    # 記錄系統事件
    if is_event_open(request) :       
        log = Log(user_id=request.user.id, event=u'下載教學筆記<'+classroom.name+u'>到Word')
        log.save()  
        
    document = Document()
    docx_title="Note-"+ classroom.name.encode("utf-8") + "-"+ str(timezone.localtime(timezone.now()).date())+".docx"

    notes = Note.objects.filter(user_id=request.user.id, classroom_id=classroom_id).order_by("-id")
    document.add_paragraph(request.user.first_name + u'的教學筆記')
    document.add_paragraph(u"班級：" + classroom.name)
    table = document.add_table(rows=1, cols=3)
    table.style = 'TableGrid'    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'課別'
    hdr_cells[1].text = u'日期'
    hdr_cells[2].text = u'內容'
    for note in notes:
        row_cells = table.add_row().cells
        row_cells[0].text = note.lesson
        row_cells[1].text = str(timezone.localtime(note.publication_date).strftime("%b %d %Y %H:%M:%S"))
        h = HTML2Text()
        h.ignore_links = True
        row_cells[2].text = h.handle(note.memo)

    # Prepare document for download        
    # -----------------------------
    f = StringIO.StringIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length


    return response
	
# 日曆：班級登入列表
class CalendarView(ListView):
    context_object_name = 'lists'
    #paginate_by = 50
    template_name = 'teacher/calendar.html'

    def get_queryset(self):    
        # 記錄系統事件
        classrooms = Classroom.objects.filter(teacher_id=self.kwargs['user_id']).order_by("-id")
        log = Log(user_id=self.request.user.id, event=u'查看登入記錄')
        log.save()
        querysets = []
        for classroom in classrooms:
            enrolls = Enroll.objects.filter(classroom_id=classroom.id).order_by("seat")
            members = []
            for enroll in enrolls:
                  if enroll.seat > 0 :
                      members.append(enroll.student_id)            
            user_logs = Log.objects.filter(user_id__in=members, event="登入系統").order_by("id")
            #weeklogs = groupby(user_logs, key=lambda row: (localtime(row.publish).isocalendar()[1]))
            logs = groupby(user_logs, key=lambda row: (localtime(row.publish).year, localtime(row.publish).month, localtime(row.publish).day))
            month_lists = []
            for key, value in logs:
                events = list(value)
                month_lists.append([key,events])
            if len(month_lists) > 0 :
                querysets.append([classroom, month_lists, (month_lists[-1][0][0]-month_lists[0][0][0])*150+200])
            else :
                querysets.append([classroom, month_lists,200])
        return querysets
        
    def get_context_data(self, **kwargs):
        context = super(CalendarView, self).get_context_data(**kwargs)			
        return context	
			
# 日曆：班級登入列表
class CalendarLogView(ListView):
    context_object_name = 'events'
    paginate_by = 50
    template_name = 'teacher/calendar_log.html'

    def get_queryset(self):    
        classroom = Classroom.objects.get(id=self.kwargs['classroom_id'])
        year = self.kwargs['year']
        month = self.kwargs['month']
        day = self.kwargs['day']
        members = []
        enrolls = Enroll.objects.filter(classroom_id=classroom.id) 
        for enroll in enrolls:
            members.append(enroll.student_id)
        # 記錄系統事件
        if is_event_open(self.request) :           
            log = Log(user_id=self.request.user.id, event=u'查看班級登入<'+year+"-"+month+"-"+day+'>')
            log.save()     
        queryset = Log.objects.filter(user_id__in=members, event="登入系統", publish__year=year, publish__month=month, publish__day=day).order_by('-id')
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(CalendarLogView, self).get_context_data(**kwargs)
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        year = self.kwargs['year']
        month = self.kwargs['month']
        day = self.kwargs['day']
        context['date'] = [year,month,day]
        return context	
			
# 影片：班級觀看列表
class VideoView(ListView):
    context_object_name = 'lists'
    #paginate_by = 50
    template_name = 'teacher/video.html'

    def get_queryset(self):    
        # 記錄系統事件
        log = Log(user_id=self.request.user.id, event=u'查看影片統計')
        log.save()
        enrolls = Enroll.objects.filter(classroom_id=self.kwargs['classroom_id'], seat__gt=0).order_by("seat")
        sids = map(lambda e: e.student_id, enrolls)
        video_pool = VideoLogHelper().getPlayLogByUserids(sids, '5', u'影片1')
        events = []
        for student in enrolls:
            videos = map(lambda v: v['start'], filter(lambda v: v['uid'] == student.student_id, video_pool))
            events.append({'seat': student.seat, 'start': videos})

        return events		
        
    def get_context_data(self, **kwargs):
        context = super(VideoView, self).get_context_data(**kwargs)			
        return context				